#!/usr/bin/env python3
"""Generate resume.pdf (and resume.docx where needed) from resume text sources.

Usage (from repo root):
    pip install fpdf2 python-docx
    python test_cases/generate_test_pdfs.py

Scans test_cases/ for resume.txt and resumes/*.txt:
  - Always writes resume.pdf (or sibling .pdf)
  - Writes resume.docx when expected_output.json sets "resume_format": "docx"
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

try:
    from fpdf import FPDF
except ImportError:
    print("Missing dependency. Install with:  pip install fpdf2", file=sys.stderr)
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore[misc, assignment]

ROOT = Path(__file__).resolve().parent
SKIP_MARKERS = ("(no resume", "(none)", "csv-only", "intentionally empty")


def text_to_pdf(text: str, output_path: Path) -> None:
    """Render plain-text resume content into a simple single-column PDF."""
    pdf = FPDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)

    stripped = text.strip()
    if not stripped or any(m in stripped.lower() for m in SKIP_MARKERS):
        pdf.cell(0, 6, " ", new_x="LMARGIN", new_y="NEXT")
    else:
        for line in text.splitlines():
            safe = line.encode("latin-1", errors="replace").decode("latin-1").rstrip()
            if not safe:
                pdf.ln(4)
                continue
            while safe:
                pdf.cell(0, 5, safe[:100], new_x="LMARGIN", new_y="NEXT")
                safe = safe[100:]

    output_path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(output_path))


def text_to_docx(text: str, output_path: Path) -> None:
    """Render plain-text resume content into a simple paragraph-per-line DOCX."""
    if Document is None:
        raise RuntimeError("Missing dependency. Install with:  pip install python-docx")

    doc = Document()
    stripped = text.strip()
    if not stripped or any(m in stripped.lower() for m in SKIP_MARKERS):
        doc.add_paragraph(" ")
    else:
        for line in text.splitlines():
            doc.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _needs_docx(tc_dir: Path) -> bool:
    expected = tc_dir / "expected_output.json"
    if not expected.exists():
        return False
    try:
        data = json.loads(expected.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return False
    return isinstance(data, dict) and data.get("resume_format") == "docx"


def main() -> None:
    txt_sources: list[Path] = []
    txt_sources.extend(sorted(ROOT.rglob("resume.txt")))
    for resumes_dir in sorted(ROOT.rglob("resumes")):
        if resumes_dir.is_dir():
            txt_sources.extend(sorted(resumes_dir.glob("*.txt")))

    if not txt_sources:
        print(f"No resume text files found under {ROOT}")
        return

    pdf_count = docx_count = 0
    for resume_txt in txt_sources:
        content = resume_txt.read_text(encoding="utf-8")
        tc_dir = resume_txt.parent
        pdf_path = resume_txt.with_suffix(".pdf")
        text_to_pdf(content, pdf_path)
        rel = resume_txt.relative_to(ROOT.parent)
        print(f"Created {rel.with_suffix('.pdf')}")
        pdf_count += 1

        if _needs_docx(tc_dir):
            docx_path = resume_txt.with_suffix(".docx")
            text_to_docx(content, docx_path)
            print(f"Created {rel.with_suffix('.docx')}")
            docx_count += 1

    print(f"\nDone — generated {pdf_count} PDF(s), {docx_count} DOCX(s).")


if __name__ == "__main__":
    main()
